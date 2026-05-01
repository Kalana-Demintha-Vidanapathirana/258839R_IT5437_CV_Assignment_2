import pathlib
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).parent.parent
OUT  = ROOT / "output"

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(text, size=8.5, bold=False, sa=3):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(0)
    return p

def code(text, size=7.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x20, 0x60, 0x20)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear"); shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F4F0")
    p._p.get_or_add_pPr().append(shading)
    return p

def img(name, width=Inches(2.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    try:
        p.add_run().add_picture(str(OUT / name), width=width)
    except Exception as e:
        p.add_run("[Image: %s]" % name)

def img_row(names, widths):
    t = doc.add_table(rows=1, cols=len(names))
    t.style = "Table Grid"
    for i, (name, w) in enumerate(zip(names, widths)):
        cp = t.rows[0].cells[i].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            cp.add_run().add_picture(str(OUT / name), width=w)
        except Exception:
            cp.add_run("[%s]" % name)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def spacer(pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(pt)
    p.paragraph_format.space_before = Pt(0)

# ── HEADER ────────────────────────────────────────────────────────
h = doc.add_heading("IT5437 — Computer Vision  |  Assignment 2", 0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h.runs:
    run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Index: 258839R     Name: Kalana Vidanapathirana")
r.font.size = Pt(10); r.font.bold = True

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("GitHub: https://github.com/Kalana-Demintha-Vidanapathirana/Assignment_2")
r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
spacer(6)

# ── Q1: LINE FITTING ──────────────────────────────────────────────
heading("1. Line Fitting")

heading("(a) Total Least Squares — Line 1", 2)
body("TLS minimises orthogonal (perpendicular) distances to the line, unlike ordinary LS which "
     "minimises vertical residuals. The line is in normal form ax+by=d. The unit normal (a,b) "
     "is the eigenvector of scatter matrix UTU for the smallest eigenvalue — last row of VT from SVD.")
code("U = pts1 - np.mean(pts1, axis=0)  # centre data\n"
     "_, _, Vt = np.linalg.svd(U.T @ U) # scatter matrix SVD\n"
     "a, b = Vt[-1, 0], Vt[-1, 1]       # normal = last row of Vt\n"
     "d    = a*mean[0] + b*mean[1]       # distance from origin")
body("Result:  a = -0.7736,  b = 0.6337,  d = -3.7942   |   Mean perpendicular residual = 0.406")
img("p1a_tls.png", width=Inches(2.3))

heading("(b) RANSAC — Three Lines", 2)
body("All 300 mixed points (3 x 100) fed to RANSAC. Each iteration: (1) sample 2 points, compute "
     "candidate line; (2) count inliers within t=0.5; (3) refit with inliers via TLS; "
     "(4) mask inliers and repeat for the next line.")
code("for i in range(3):\n"
     "    model, mask = ransac_line(pts_all[remaining], t=0.5, max_iter=500)\n"
     "    remaining[global_idx[mask]] = False")
body("Line 1: 0.442x+0.897y=1.881 (83 inliers)   Line 2: -0.719x+0.695y=0.727 (66 inliers)   "
     "Line 3: -0.787x+0.617y=-3.783 (62 inliers).  89 points unassigned near intersections.")
img("p1b_ransac.png", width=Inches(2.3))

doc.add_page_break()

# ── Q2: EARRING SIZE ──────────────────────────────────────────────
heading("2. Physical Size of Earrings")
body("Camera perpendicular to earring plane at u=720 mm. Thin lens (1/f=1/u+1/v) with f=8 mm "
     "gives v=8.090 mm, magnification m=0.01124. Scale = pixel_size/m = 0.1958 mm/pixel.")
code("v = (f*u)/(u-f)   # 8.090 mm\n"
     "m = v/u            # 0.01124\n"
     "mm_per_px = p/m    # 0.1958 mm/px\n"
     "# HSV threshold -> RETR_CCOMP contours -> minEnclosingCircle")
body("Both earrings identical.  Outer: 418 px -> 81.84 mm.  "
     "Inner: 316 px -> 61.87 mm.  Band: 51 px -> 9.99 mm.")
img("p2_earring_sizes.png", width=Inches(3.5))

doc.add_page_break()

# ── Q3: CIRCUIT BOARD ALIGNMENT ───────────────────────────────────
heading("3. Circuit Board Alignment")
body("c1.jpg and c2.jpg are the same Arduino Mega board from different viewpoints (manually transformed).")

heading("(a) Manual Homography", 2)
body("6 landmarks clicked via OpenCV mouse-callback (IMREAD_REDUCED_COLOR_4 = 1/4 res). "
     "Coordinates scaled x4 to full resolution. Homography via DLT (no RANSAC).")
code("H, _ = cv.findHomography(pts_c1_full, pts_c2_full, method=0)\n"
     "c1_warped = cv.warpPerspective(c1, H, (w2, h2))")
img("p3a_warp.png", width=Inches(5.5))

heading("(b) Image Difference", 2)
body("Same board -> ideal difference is zero. Residual bright pixels reflect manual alignment "
     "error or lighting change, not actual board differences.")
img("p3b_diff.png", width=Inches(5.0))

doc.add_page_break()

heading("(c) SIFT Feature Matching", 2)
body("SIFT: 4172/4500 keypoints on c1/c2 (0.4x scale). FLANN + Lowe ratio test (t=0.7) -> "
     "2641 good matches. Consistent match directions confirm a coherent perspective transform.")
code("sift = cv.SIFT_create()\n"
     "kp1, des1 = sift.detectAndCompute(g1, None)\n"
     "good = [m for m,n in flann.knnMatch(des1,des2,k=2) if m.distance < 0.7*n.distance]")
img("p3c_sift_matches.png", width=Inches(5.5))

heading("(d) Automatic Homography + Comparison", 2)
body("H_auto via cv.findHomography with RANSAC (reprojThreshold=5px), 2523/2641 inlier SIFT matches.")
code("H_auto, _ = cv.findHomography(src_pts, dst_pts, cv.RANSAC, 5.0)\n"
     "# H_auto ~ [[0.9645,-0.2638,406.2],[0.2637,0.9645,-384.0],[0,0,1]]")
img_row(["p3d_warp_comparison.png", "p3d_diff_comparison.png"], [Inches(2.9), Inches(2.7)])

heading("Discussion: Manual vs Automatic Homography", 2)
tbl = doc.add_table(rows=5, cols=3); tbl.style = "Table Grid"
for j, h in enumerate(["", "Manual H (a)", "Auto SIFT H (d)"]):
    r = tbl.rows[0].cells[j].paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(8)
for i, row in enumerate([
    ("Correspondences",  "6 clicked points", "2523 RANSAC inliers"),
    ("Outlier rejection","None (DLT only)",   "RANSAC built-in"),
    ("Human error",      "Yes (~5-10 px)",    "None"),
    ("Mean pixel diff",  "Larger",            "~0.8/255 (near-perfect)"),
], 1):
    for j, val in enumerate(row):
        tbl.rows[i].cells[j].paragraphs[0].add_run(val).font.size = Pt(8)
spacer(4)
body("1. Auto H far more accurate: 2523 constraints vs 6 -> better-conditioned solve.\n"
     "2. Manual H sensitive to click precision; 5 px at 1/4-res = 20 px at full res.\n"
     "3. RANSAC rejects ~4% outlier matches automatically, making H_auto robust.\n"
     "4. Both matrices agree on gross transform (~15 deg rotation, ~400 px translation).\n"
     "5. Near-zero auto diff confirms same board; SIFT homography achieves near-perfect alignment.")

doc.save(str(ROOT / "258839R.docx"))
print("Saved", ROOT / "258839R.docx")
